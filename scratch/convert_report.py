import os
import re
import base64
import tempfile
import lxml.html
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_element(name):
    return OxmlElement(name)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = parse_xml(r'<w:tblBorders %s>'
                        r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                        r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                        r'  <w:left w:val="none"/>'
                        r'  <w:right w:val="none"/>'
                        r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E5E5E5"/>'
                        r'  <w:insideV w:val="none"/>'
                        r'</w:tblBorders>' % nsdecls('w'))
    tblPr.append(borders)

def add_page_number(run):
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    run._r.append(fldSimple)

def add_header_footer(section, fmt='decimal'):
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False
    
    # Header (Right-aligned, gray, italic)
    header_para = section.header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header_run = header_para.add_run("99Care OS – AI-Powered Healthcare Operations Management System")
    header_run.font.name = 'Times New Roman'
    header_run.font.size = Pt(8.5)
    header_run.font.italic = True
    header_run.font.color.rgb = RGBColor(120, 120, 120)
    
    # Footer (Centered page number)
    footer_para = section.footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run("Auro University, Surat  |  ")
    footer_run.font.name = 'Times New Roman'
    footer_run.font.size = Pt(9.5)
    footer_run.font.color.rgb = RGBColor(80, 80, 80)
    
    page_run = footer_para.add_run("Page ")
    page_run.font.name = 'Times New Roman'
    page_run.font.size = Pt(9.5)
    page_run.font.color.rgb = RGBColor(80, 80, 80)
    
    add_page_number(page_run)

def add_paragraph_with_runs(p, elem):
    if elem.text:
        p.add_run(elem.text)
    for child in elem.getchildren():
        if child.tag == 'br':
            p.add_run('\n')
        elif child.tag in ['strong', 'b']:
            run = p.add_run(child.text_content())
            run.font.bold = True
        elif child.tag in ['em', 'i']:
            run = p.add_run(child.text_content())
            run.font.italic = True
        else:
            p.add_run(child.text_content())
        if child.tail:
            p.add_run(child.tail)

def process_table(table_node, doc):
    rows = table_node.xpath(".//tr")
    if not rows:
        return
    
    # Count max columns
    max_cols = 0
    for r in rows:
        cols = r.xpath("./td | ./th")
        max_cols = max(max_cols, len(cols))
    
    if max_cols == 0:
        return
        
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    for r_idx, r in enumerate(rows):
        cells = r.xpath("./td | ./th")
        for c_idx, c in enumerate(cells):
            if c_idx >= max_cols:
                continue
            cell = table.cell(r_idx, c_idx)
            
            # Style formatting of cells
            p = cell.paragraphs[0]
            p.style = 'Normal'
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            
            add_paragraph_with_runs(p, c)
            
            # Set cell margins for neat look
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            
            # If table header
            if c.tag == 'th' or 'col-bold' in c.get('class', ''):
                if p.runs:
                    for run in p.runs:
                        run.font.bold = True

def process_list(elem, doc):
    tag = elem.tag
    items = elem.xpath(".//li")
    for item in items:
        p = doc.add_paragraph(style='List Bullet' if tag == 'ul' else 'List Number')
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)
        add_paragraph_with_runs(p, item)

def process_image(img_node, doc, html_path):
    src = img_node.get('src', '')
    if src.startswith('data:image/png;base64,'):
        # Decode base64 image
        img_data = base64.b64decode(src.split(',')[1])
        with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as temp_img:
            temp_img.write(img_data)
            temp_path = temp_img.name
        try:
            # Insert picture centered
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(temp_path, width=Inches(5.8))
            os.unlink(temp_path)
        except Exception as e:
            print(f"Error inserting image: {e}")
    elif src:
        # Check relative file
        possible_paths = [
            os.path.join(os.path.dirname(html_path), src),
            os.path.join("/Users/tanishqkachiwala/Downloads/Design/healthcare", src),
            os.path.join("/Users/tanishqkachiwala/Downloads/Design/healthcare/report", os.path.basename(src))
        ]
        inserted = False
        for path in possible_paths:
            if os.path.exists(path):
                try:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(path, width=Inches(5.8))
                    inserted = True
                    print(f"Inserted image from path: {path}")
                    break
                except Exception as e:
                    print(f"Error inserting picture from {path}: {e}")
        if not inserted:
            print(f"Warning: Image file not found or failed to load: {src}")

def process_element(elem, doc, html_path):
    tag = elem.tag
    cls = elem.get('class', '') or ''
    
    # Skip headers/footers/page borders/logo-header
    if 'pg-header' in cls or 'pg-footer' in cls or 'page-border' in cls or 'logo-header' in cls:
        return
        
    # Check if table
    if tag == 'table' or 'tbl-block' in cls:
        table_node = elem if tag == 'table' else elem.xpath(".//table")
        if table_node:
            process_table(table_node[0], doc)
        return
        
    # Check if list
    if tag in ['ul', 'ol'] or 'bullet-list' in cls or 'num-list' in cls:
        process_list(elem, doc)
        return
        
    # Check if image block
    if tag == 'img' or 'img-block' in cls or 'img-center' in cls:
        img_nodes = [elem] if tag == 'img' else elem.xpath(".//img")
        for img in img_nodes:
            process_image(img, doc, html_path)
        return

    # Check if chapter title/major heading
    if 'ch-title' in cls or 'doc-title' in cls:
        text = elem.text_content().strip().upper()
        if text:
            doc.add_paragraph(text, style='ChapterTitle')
        return

    # Check if heading 1
    if 'sec-head' in cls or tag == 'h1':
        text = elem.text_content().strip()
        if text:
            doc.add_paragraph(text, style='Heading 1')
        return

    # Check if heading 2
    if 'sub-head' in cls or tag == 'h2':
        text = elem.text_content().strip()
        if text:
            doc.add_paragraph(text, style='Heading 2')
        return

    # Check if title line on the cover page
    if 'title-line' in cls:
        text = elem.text_content().strip()
        if text:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.style = 'Normal'
            # Remove high paragraph spacing for cover page spacing control
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            
            run = p.add_run(text)
            run.font.name = 'Times New Roman'
            if 'small' in cls:
                run.font.size = Pt(11)
            else:
                run.font.size = Pt(14)
                run.font.bold = True
        return

    # Check if standard paragraph
    if tag == 'p' or 'body-p' in cls or 'doc-p' in cls or 'doc-field' in cls or 'sig-block' in cls or 'tbl-caption' in cls or 'img-caption' in cls or 'fig-caption' in cls:
        # Use our inline formatter to preserve bold/italic/br tags
        p = doc.add_paragraph()
        p.style = 'Normal'
        
        if 'sig-block' in cls:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            add_paragraph_with_runs(p, elem)
            if p.runs:
                for r in p.runs:
                    r.font.bold = True
        elif 'tbl-caption' in cls or 'img-caption' in cls or 'fig-caption' in cls:
            # If it's a figure caption, dynamically insert the visual PNG image right before the caption paragraph!
            text = elem.text_content().strip()
            if 'Figure 7.1a' in text:
                img_path = "/Users/tanishqkachiwala/Downloads/Design/healthcare/report/flowcharts/7_1_system_flowchart.png"
                if os.path.exists(img_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(img_path, width=Inches(6.0))
            elif 'Figure 7.4' in text:
                img_path = "/Users/tanishqkachiwala/Downloads/Design/healthcare/report/flowcharts/7_4_erd.png"
                if os.path.exists(img_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(img_path, width=Inches(6.0))
            elif 'Figure 7.6' in text:
                img_path = "/Users/tanishqkachiwala/Downloads/Design/healthcare/report/flowcharts/7_6_class_diagram.png"
                if os.path.exists(img_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(img_path, width=Inches(5.0))
            elif 'Figure 7.7' in text:
                img_path = "/Users/tanishqkachiwala/Downloads/Design/healthcare/report/flowcharts/7_7_lead_intake_sequence.png"
                if os.path.exists(img_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(img_path, width=Inches(5.0))
            elif 'Figure 7.8' in text:
                img_path = "/Users/tanishqkachiwala/Downloads/Design/healthcare/report/flowcharts/7_8_worker_assignment_sequence.png"
                if os.path.exists(img_path):
                    p_img = doc.add_paragraph()
                    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_img.add_run().add_picture(img_path, width=Inches(5.0))
            
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_after = Pt(6)
            add_paragraph_with_runs(p, elem)
            if p.runs:
                for r in p.runs:
                    r.font.bold = True
                    r.font.size = Pt(11)
        else:
            add_paragraph_with_runs(p, elem)
        return

    # If it's a generic div or other container, descend to its children
    for child in elem.getchildren():
        process_element(child, doc, html_path)

def main():
    html_path = "/Users/tanishqkachiwala/Downloads/Design/tech-h0use/99care_full_report.html"
    docx_path = "/Users/tanishqkachiwala/Downloads/Design/tech-h0use/99care_full_report.docx"
    
    print(f"Reading HTML file: {html_path}")
    with open(html_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    parser = lxml.html.HTMLParser(encoding='utf-8')
    doc_tree = lxml.html.fromstring(html_content.encode('utf-8'), parser=parser)
    
    doc = Document()
    
    # Global styles setup
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.paragraph_format.line_spacing = 2.0  # Double-spaced
    style_normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Custom styles
    style_ch = doc.styles.add_style('ChapterTitle', 1)
    style_ch.font.name = 'Times New Roman'
    style_ch.font.size = Pt(14)
    style_ch.font.bold = True
    style_ch.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_ch.paragraph_format.line_spacing = 1.5
    style_ch.paragraph_format.space_after = Pt(16)
    
    style_h1 = doc.styles['Heading 1']
    style_h1.font.name = 'Times New Roman'
    style_h1.font.size = Pt(12)
    style_h1.font.bold = True
    style_h1.font.color.rgb = RGBColor(0, 0, 0)
    style_h1.paragraph_format.line_spacing = 2.0
    style_h1.paragraph_format.space_before = Pt(20)
    style_h1.paragraph_format.space_after = Pt(6)
    
    style_h2 = doc.styles['Heading 2']
    style_h2.font.name = 'Times New Roman'
    style_h2.font.size = Pt(12)
    style_h2.font.bold = True
    style_h2.font.underline = True
    style_h2.font.color.rgb = RGBColor(0, 0, 0)
    style_h2.paragraph_format.line_spacing = 2.0
    style_h2.paragraph_format.space_before = Pt(14)
    style_h2.paragraph_format.space_after = Pt(4)
    
    # Setup first section margins
    first_sec = doc.sections[0]
    first_sec.page_width = Inches(8.27)
    first_sec.page_height = Inches(11.69)
    first_sec.top_margin = Inches(1)
    first_sec.right_margin = Inches(1)
    first_sec.bottom_margin = Inches(1)
    first_sec.left_margin = Inches(1.5)
    
    # Track section types
    # 0 = Cover/Certificates (no header/footer)
    # 1 = Prelims (Roman page numbering i, ii, iii...)
    # 2 = Main (Arabic page numbering 1, 2, 3...)
    current_section_type = 0
    current_section = first_sec
    
    pages = doc_tree.xpath("//div[contains(@class, 'screen-page')]")
    print(f"Found {len(pages)} pages in the report HTML.")
    
    for i, page in enumerate(pages):
        # Extract the page body/content container
        body = page.xpath(".//div[@class='page-body'] | .//div[contains(@class, 'page-content')] | .//div[@class='page-inner']")
        if not body:
            body = page
        else:
            body = body[0]
            
        # Check for section transitions
        section_added = False
        ch_titles = body.xpath(".//div[@class='ch-title']")
        if ch_titles:
            title_text = ch_titles[0].text_content().strip().upper()
            if "ACKNOWLEDGEMENT" in title_text and current_section_type < 1:
                # Add section break for preliminary pages (Roman numerals)
                prelim_sec = doc.add_section()
                prelim_sec.page_width = Inches(8.27)
                prelim_sec.page_height = Inches(11.69)
                prelim_sec.top_margin = Inches(1)
                prelim_sec.right_margin = Inches(1)
                prelim_sec.bottom_margin = Inches(1)
                prelim_sec.left_margin = Inches(1.5)
                
                prelim_sec_pr = prelim_sec._sectPr
                pgNumType = OxmlElement('w:pgNumType')
                pgNumType.set(qn('w:fmt'), 'lrRoman')
                pgNumType.set(qn('w:start'), '1') # Starts at i
                prelim_sec_pr.append(pgNumType)
                
                current_section = prelim_sec
                current_section_type = 1
                section_added = True
                add_header_footer(current_section, fmt='lrRoman')
                print(f"Started Preliminary Section (Roman i) at page {i+1}.")
                
            elif "CHAPTER 2" in title_text and current_section_type < 2:
                # Add section break for main chapters (Arabic numerals starting from 1)
                main_sec = doc.add_section()
                main_sec.page_width = Inches(8.27)
                main_sec.page_height = Inches(11.69)
                main_sec.top_margin = Inches(1)
                main_sec.right_margin = Inches(1)
                main_sec.bottom_margin = Inches(1)
                main_sec.left_margin = Inches(1.5)
                
                main_sec_pr = main_sec._sectPr
                pgNumType = OxmlElement('w:pgNumType')
                pgNumType.set(qn('w:fmt'), 'decimal')
                pgNumType.set(qn('w:start'), '1') # Starts at 1
                main_sec_pr.append(pgNumType)
                
                current_section = main_sec
                current_section_type = 2
                section_added = True
                add_header_footer(current_section, fmt='decimal')
                print(f"Started Main Section (Arabic 1) at page {i+1}.")
                
        if not section_added and i > 0:
            doc.add_page_break()
            
        # Process elements recursively
        for elem in body.getchildren():
            process_element(elem, doc, html_path)
            
    print(f"Saving DOCX report to: {docx_path}")
    doc.save(docx_path)
    print("DOCX report saved successfully!")

if __name__ == "__main__":
    main()
